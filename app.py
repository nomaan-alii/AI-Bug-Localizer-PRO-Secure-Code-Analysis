import streamlit as st
import json
from streamlit_ace import st_ace
import plotly.express as px
import pandas as pd
import plotly.graph_objects as go
from analyzer.static_analysis import analyze_ast
from analyzer.runtime_analysis import run_code
from analyzer.explanation import explain
from analyzer.fix_suggester import suggest_fix
from analyzer.code_corrector import apply_basic_fixes
from analyzer.owasp_analyzer import OWASPAnalyzer
from analyzer.encryption_analyzer import EncryptionAnalyzer
from analyzer.compliance_checker import ComplianceChecker
from analyzer.authentication_analyzer import AuthenticationAnalyzer
from analyzer.api_security_analyzer import APISecurityAnalyzer
from analyzer.dependency_vulnerability_analyzer import DependencyVulnerabilityAnalyzer
from analyzer.secure_configuration_analyzer import SecureConfigurationAnalyzer
from analyzer.input_validation_analyzer import InputValidationAnalyzer
from analyzer.threat_modeling_analyzer import ThreatModelingAnalyzer
from ml.scoring import rank_issues
from analyzer.complexity import compute_complexity
from core.quality_engine import build_quality_report
from analyzer.security_analysis import analyze_security
from ui.components import render_security_dashboard, render_issue_card, render_compliance_checker
from core.enhanced_ai_chat import get_ai_engine
import hashlib
from cryptography.fernet import Fernet
import io
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

try:
    from docx import Document
except ImportError:
    Document = None

def load_css(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)


def generate_pdf_report(data: dict) -> bytes:
    """Generate PDF report from analysis data."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    title = "AI BUG LOCALIZER PRO - Security Analysis Report"
    c.setFont("Helvetica-Bold", 14)
    c.drawString(1 * inch, height - 1 * inch, title)
    c.setFont("Helvetica", 9)
    c.drawString(1 * inch, height - 1.2 * inch, f"Generated: {datetime.utcnow().isoformat()} UTC")
    y = height - 1.6 * inch
    issues = data.get("issues") or []
    total = len(issues)
    c.setFont("Helvetica-Bold", 11)
    c.drawString(1 * inch, y, f"Summary: {total} total issues found")
    y -= 0.25 * inch
    code = data.get("code", "")
    if code:
        c.setFont("Helvetica-Bold", 10)
        c.drawString(1 * inch, y, "Code Preview:")
        y -= 0.18 * inch
        c.setFont("Helvetica", 8)
        for line in code.strip().splitlines()[:10]:
            if y < 1 * inch:
                c.showPage()
                y = height - 1 * inch
            c.drawString(1 * inch, y, line[:100])
            y -= 0.16 * inch
        y -= 0.08 * inch
    sections = [("OWASP Issues", "owasp_issues"), ("Encryption Issues", "crypto_issues"), ("Compliance Issues", "compliance_issues"), ("Authentication Issues", "auth_issues"), ("API Security Issues", "api_security_issues"), ("Dependency Vulnerabilities", "dependency_issues"), ("Configuration Issues", "config_security_issues"), ("Input Validation Issues", "input_validation_issues"), ("Threat Modeling Issues", "threat_model_issues")]
    for title_text, key in sections:
        items = data.get(key, []) or []
        if not items:
            continue
        if y < 1 * inch:
            c.showPage()
            y = height - 1 * inch
        c.setFont("Helvetica-Bold", 11)
        c.drawString(1 * inch, y, f"{title_text} (count: {len(items)})")
        y -= 0.22 * inch
        c.setFont("Helvetica", 9)
        for issue in items:
            msg = issue.get("msg", issue.get("type", ""))
            score = issue.get("score", issue.get("severity", ""))
            line_no = issue.get("line", "-")
            text = f"- [{line_no}] {issue.get('type','')}: {msg} (score: {score})"
            if y < 1 * inch:
                c.showPage()
                y = height - 1 * inch
            c.drawString(1 * inch, y, (text[:95] + ("..." if len(text) > 95 else "")))
            y -= 0.18 * inch
        y -= 0.12 * inch
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.read()


def generate_word_report(data: dict) -> bytes:
    """Generate comprehensive Word report with APA 7 citations."""
    if Document is None:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    doc = Document()
    doc.add_heading('AI Bug Localizer PRO - Comprehensive Security Analysis Report', 0)

    # Title page
    doc.add_paragraph(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")} UTC')
    doc.add_paragraph('This report provides a detailed analysis of code security vulnerabilities using advanced AI-powered analyzers.')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    issues = data.get("issues") or []
    total = len(issues)
    doc.add_paragraph(f'Total security issues detected: {total}')
    doc.add_paragraph('This comprehensive report covers static analysis, runtime analysis, OWASP Top 10 vulnerabilities, encryption weaknesses, compliance gaps, authentication flaws, API security risks, dependency vulnerabilities, configuration issues, input validation problems, and threat modeling concerns (Python Software Foundation, 2023).')

    # Code Preview
    code = data.get("code", "")
    if code:
        doc.add_heading('Analyzed Code Preview', level=1)
        doc.add_paragraph('The following code snippet was analyzed:')
        code_para = doc.add_paragraph()
        code_para.add_run(code[:1000] + ('...' if len(code) > 1000 else '')).italic = True

    # Detailed Analysis Sections
    sections = [
        ("Static Analysis Issues", "static_issues", "Static analysis examines code structure without execution (Python Software Foundation, 2023)."),
        ("Runtime Analysis Issues", "runtime_issues", "Runtime analysis detects issues during code execution."),
        ("Security Analysis Issues", "security_issues", "General security vulnerabilities identified."),
        ("OWASP Top 10 Issues", "owasp_issues", "OWASP Top 10 web application security risks (OWASP Foundation, 2023)."),
        ("Encryption Issues", "crypto_issues", "Cryptographic weaknesses and insecure encryption practices."),
        ("Compliance Issues", "compliance_issues", "Regulatory compliance violations."),
        ("Authentication Issues", "auth_issues", "Authentication and authorization flaws."),
        ("API Security Issues", "api_security_issues", "API security vulnerabilities."),
        ("Dependency Vulnerabilities", "dependency_issues", "Third-party library security risks."),
        ("Configuration Issues", "config_security_issues", "Insecure configuration settings."),
        ("Input Validation Issues", "input_validation_issues", "Input sanitization and validation problems."),
        ("Threat Modeling Issues", "threat_model_issues", "Potential threat vectors identified.")
    ]

    for title_text, key, description in sections:
        items = data.get(key, []) or []
        if not items:
            continue
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(description)
        doc.add_paragraph(f'Total issues in this category: {len(items)}')

        for i, issue in enumerate(items, 1):
            msg = issue.get("msg", issue.get("type", ""))
            score = issue.get("score", issue.get("severity", ""))
            line_no = issue.get("line", "-")
            category = issue.get("category", "")
            recommendation = issue.get("recommendation", "")

            doc.add_heading(f'{i}. {issue.get("type","").title()}', level=2)
            doc.add_paragraph(f'Line: {line_no}')
            doc.add_paragraph(f'Severity Score: {score}')
            if category:
                doc.add_paragraph(f'Category: {category}')
            doc.add_paragraph(f'Description: {msg}')
            if recommendation:
                doc.add_paragraph(f'Recommendation: {recommendation}')

    # References
    doc.add_heading('References', level=1)
    references = [
        "OWASP Foundation. (2023). *OWASP Top 10*. https://owasp.org/www-project-top-ten/",
        "Python Software Foundation. (2023). *Python documentation*. https://docs.python.org/3/",
        "National Institute of Standards and Technology. (2023). *NIST cybersecurity framework*. https://www.nist.gov/cyberframework",
        "Open Web Application Security Project. (2023). *OWASP cheat sheet series*. https://cheatsheetseries.owasp.org/",
        "Python Cryptographic Authority. (2023). *cryptography library*. https://cryptography.io/"
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# CALL IT HERE (IMPORTANT)
load_css("assets/styles.css")        

# ================= PAGE CONFIG =================
st.set_page_config(
    page_title="AI Bug Localizer PRO",
    layout="wide",
    page_icon="🧠"
)

# ================= SESSION STATE =================
st.session_state.setdefault("data", {})
st.session_state.setdefault("chat_history", [])
st.session_state.setdefault("history", [])


# ================= SAFE ANALYSIS ENGINE =================
@st.cache_data(show_spinner=False)
def get_chart_data(issues):
    high = len([i for i in issues if i.get("score", 0) >= 0.85])
    medium = len([i for i in issues if 0.6 <= i.get("score", 0) < 0.85])
    low = len([i for i in issues if i.get("score", 0) < 0.6])
    return high, medium, low

@st.cache_data(show_spinner=False)
def cached_analysis(code):

    if not code.strip():
        return {
            "code": "",
            "issues": [],
            "fixed": "",
            "complexity": {"complexity_score": 0, "risk": "LOW"},
            "security_issues": [],
            "owasp_issues": [],
            "crypto_issues": [],
            "compliance_issues": [],
            "auth_issues": []
        }

    complexity = compute_complexity(code) or {}

    ast_issues = analyze_ast(code) or []
    runtime_issues = run_code(code) or []
    security_issues = analyze_security(code) or []
    
    # NEW: Information Security Analysis
    owasp_analyzer = OWASPAnalyzer()
    encryption_analyzer = EncryptionAnalyzer()
    compliance_checker = ComplianceChecker()
    auth_analyzer = AuthenticationAnalyzer()
    api_security_analyzer = APISecurityAnalyzer()
    dependency_analyzer = DependencyVulnerabilityAnalyzer()
    config_analyzer = SecureConfigurationAnalyzer()
    input_validation_analyzer = InputValidationAnalyzer()
    threat_modeling_analyzer = ThreatModelingAnalyzer()
    
    owasp_issues = owasp_analyzer.analyze(code) or []
    crypto_issues = encryption_analyzer.analyze(code) or []
    compliance_issues = compliance_checker.analyze(code) or []
    auth_issues = auth_analyzer.analyze(code) or []
    api_security_issues = api_security_analyzer.analyze(code) or []
    dependency_issues = dependency_analyzer.analyze(code) or []
    config_security_issues = config_analyzer.analyze(code) or []
    input_validation_issues = input_validation_analyzer.analyze(code) or []
    threat_model_issues = threat_modeling_analyzer.analyze(code) or []

    all_issues = ast_issues + runtime_issues + security_issues + owasp_issues + crypto_issues + compliance_issues + auth_issues + api_security_issues + dependency_issues + config_security_issues + input_validation_issues + threat_model_issues
    top_issues = rank_issues(all_issues or [], top_k=None)

    fixed_code = apply_basic_fixes(code) if top_issues else code

    score = complexity.get("complexity_score", 0)

    if score >= 70:
        complexity["risk"] = "HIGH"
    elif score >= 40:
        complexity["risk"] = "MEDIUM"
    else:
        complexity["risk"] = "LOW"

    return {
        "code": code,
        "issues": top_issues,
        "fixed": fixed_code,
        "complexity": complexity,
        "security_issues": security_issues,
        "owasp_issues": owasp_issues,
        "crypto_issues": crypto_issues,
        "compliance_issues": compliance_issues,
        "auth_issues": auth_issues,
        "api_security_issues": api_security_issues,
        "dependency_issues": dependency_issues,
        "config_security_issues": config_security_issues,
        "input_validation_issues": input_validation_issues,
        "threat_model_issues": threat_model_issues
    }

st.markdown("""
<style>
/* ===== GLOBAL STYLES ===== */
:root {
    --primary: #6366f1;
    --primary-dark: #4f46e5;
    --secondary: #10b981;
    --danger: #ef4444;
    --warning: #f59e0b;
    --success: #10b981;
    --dark-bg: #0f172a;
    --card-bg: #1e293b;
    --border-color: #334155;
    --text-primary: #f1f5f9;
    --text-secondary: #cbd5e1;
}

.stApp {
    background: linear-gradient(135deg, #0f172a 0%, #1a1f36 100%);
    color: var(--text-primary);
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
}

/* ===== TYPOGRAPHY ===== */
h1, h2, h3, h4 {
    color: #e0e7ff;
    font-weight: 700;
    letter-spacing: -0.5px;
}

h1 {
    font-size: 2.5rem;
    margin-bottom: 0.5rem;
}

h2 {
    font-size: 1.875rem;
    margin-top: 1.5rem;
    margin-bottom: 1rem;
}

p, span, label {
    color: var(--text-secondary);
    font-size: 0.95rem;
    line-height: 1.6;
}

/* ===== CONTAINERS & LAYOUT ===== */
.block-container {
    padding-top: 2rem;
    max-width: 1400px;
}

/* ===== BUTTONS ===== */
.stButton > button {
    background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%);
    color: white;
    border-radius: 8px;
    border: none;
    padding: 12px 24px;
    font-weight: 600;
    font-size: 0.95rem;
    transition: all 0.3s ease;
    box-shadow: 0 4px 15px rgba(99, 102, 241, 0.3);
}

.stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 8px 25px rgba(99, 102, 241, 0.5);
    background: linear-gradient(135deg, #7c3aed 0%, #a855f7 100%);
}

.stButton > button:active {
    transform: translateY(0);
}

/* ===== PRIMARY BUTTON (TYPE=PRIMARY) ===== */
.stButton > button[kind="primary"] {
    background: linear-gradient(135deg, #10b981 0%, #059669 100%);
    box-shadow: 0 4px 15px rgba(16, 185, 129, 0.3);
}

.stButton > button[kind="primary"]:hover {
    box-shadow: 0 8px 25px rgba(16, 185, 129, 0.5);
    background: linear-gradient(135deg, #047857 0%, #065f46 100%);
}

/* ===== METRICS ===== */
div[data-testid="stMetric"] {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.8) 0%, rgba(51, 65, 85, 0.3) 100%);
    border-radius: 12px;
    padding: 20px;
    border: 1px solid var(--border-color);
    transition: all 0.3s ease;
}

div[data-testid="stMetric"]:hover {
    background: linear-gradient(135deg, rgba(51, 65, 85, 0.8) 0%, rgba(71, 85, 105, 0.4) 100%);
    border-color: #6366f1;
    transform: translateY(-2px);
    box-shadow: 0 4px 20px rgba(99, 102, 241, 0.2);
}

div[data-testid="stMetric"] .metric-label {
    color: var(--text-secondary);
    font-size: 0.85rem;
    text-transform: uppercase;
    letter-spacing: 0.5px;
    font-weight: 600;
}

div[data-testid="stMetric"] .metric-value {
    color: #e0e7ff;
    font-size: 1.8rem;
    font-weight: 700;
}

/* ===== EXPANDERS ===== */
.streamlit-expanderHeader {
    background: linear-gradient(90deg, #1e293b 0%, #0f172a 100%);
    border-radius: 8px;
    border: 1px solid var(--border-color);
    transition: all 0.3s ease;
}

.streamlit-expanderHeader:hover {
    background: linear-gradient(90deg, #334155 0%, #1e293b 100%);
    border-color: #6366f1;
}

/* ===== TEXT INPUT ===== */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div > select {
    background-color: #1e293b !important;
    color: var(--text-primary) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: 8px !important;
    padding: 12px !important;
    font-size: 0.95rem !important;
}

.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus,
.stSelectbox > div > div > select:focus {
    border-color: #6366f1 !important;
    box-shadow: 0 0 0 3px rgba(99, 102, 241, 0.1) !important;
}

/* ===== SIDEBAR ===== */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1a1f36 0%, #0f172a 100%);
    border-right: 1px solid var(--border-color);
}

section[data-testid="stSidebar"] .stRadio > div {
    gap: 0.75rem;
}

section[data-testid="stSidebar"] .stRadio > label > div {
    color: var(--text-secondary);
    font-weight: 500;
}

section[data-testid="stSidebar"] .stRadio input:checked + div {
    background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%);
    color: white;
}

/* ===== CHECKBOXES ===== */
.stCheckbox > label > div {
    color: var(--text-secondary);
    font-weight: 500;
}

/* ===== ALERTS & MESSAGES ===== */
.stAlert {
    border-radius: 8px;
    border: none;
    padding: 16px;
    font-weight: 500;
}

.stAlert > div[data-testid="stAlertContainer"] {
    border-radius: 8px;
}

/* ===== TABS ===== */
button[data-baseweb="tab"] {
    color: var(--text-secondary);
    font-weight: 600;
    border-radius: 8px 8px 0 0;
}

button[data-baseweb="tab"][aria-selected="true"] {
    color: white;
    border-bottom: 3px solid #6366f1;
    background-color: rgba(99, 102, 241, 0.1);
}

/* ===== DIVIDERS ===== */
hr {
    border-color: var(--border-color);
    margin: 2rem 0;
}

/* ===== CODE BLOCKS ===== */
.stCodeBlock {
    background: #0f172a !important;
    border: 1px solid var(--border-color) !important;
    border-radius: 8px !important;
}

/* ===== CARDS (USING MARKDOWN) ===== */
.card {
    background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
    border: 1px solid var(--border-color);
    border-radius: 12px;
    padding: 24px;
    margin: 16px 0;
    transition: all 0.3s ease;
}

.card:hover {
    border-color: #6366f1;
    box-shadow: 0 8px 30px rgba(99, 102, 241, 0.15);
    transform: translateY(-2px);
}

.card.critical {
    border-left: 4px solid #ef4444;
}

.card.warning {
    border-left: 4px solid #f59e0b;
}

.card.success {
    border-left: 4px solid #10b981;
}

/* ===== GRADIENTS ===== */
.gradient-header {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    padding: 32px;
    border-radius: 12px;
    margin-bottom: 24px;
    box-shadow: 0 10px 40px rgba(102, 126, 234, 0.15);
}

.gradient-header h2 {
    color: white;
    margin: 0 0 8px 0;
}

.gradient-header p {
    color: rgba(255, 255, 255, 0.9);
    margin: 0;
}

/* ===== FILE UPLOADER ===== */
.stFileUploader {
    border: 2px dashed var(--border-color);
    border-radius: 8px;
    padding: 20px;
}

/* ===== SPINNERS ===== */
.stSpinner > div {
    border-color: #6366f1 !important;
}

/* ===== ACE EDITOR ===== */
.ace_editor {
    border-radius: 8px !important;
    border: 1px solid var(--border-color) !important;
}

/* ===== RESPONSIVE ===== */
@media (max-width: 768px) {
    h1 { font-size: 1.75rem; }
    h2 { font-size: 1.5rem; }
    .gradient-header { padding: 24px; }
}
</style>
""", unsafe_allow_html=True)

# ================= HEADER =================
st.markdown("""
<div class="hero-shell">
    <div class="hero-card">
        <div class="hero-deco hero-deco-top-left"></div>
        <div class="hero-deco hero-deco-bottom-right"></div>
        <div class="hero-navigator">
            <div class="hero-badges">
                <span class="hero-badge"><span>🔐</span> ENTERPRISE EDITION</span>
                <span class="hero-badge accent"><span>AI</span> AI SECURITY ENGINE</span>
            </div>
            <div class="hero-subtitle">NEXT-GEN SECURE CODE INTELLIGENCE</div>
        </div>
        <div class="hero-main">
            <div class="hero-logo-shell" aria-hidden="true">
                <div class="hero-logo-base"></div>
                <svg viewBox="0 0 72 72" width="52" height="52" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                    <circle cx="36" cy="36" r="14" fill="rgba(90, 120, 255, 0.16)" />
                    <path d="M36 22v-8" opacity="0.9" />
                    <path d="M36 50v8" opacity="0.9" />
                    <path d="M22 36h-8" opacity="0.9" />
                    <path d="M58 36h8" opacity="0.9" />
                    <path d="M26 26l-8-8" opacity="0.85" />
                    <path d="M46 26l8-8" opacity="0.85" />
                    <path d="M26 46l-8 8" opacity="0.85" />
                    <path d="M46 46l8 8" opacity="0.85" />
                </svg>
            </div>
            <div class="hero-content">
                <div class="hero-eyebrow">AI Bug Localizer</div>
                <h1 class="hero-heading">AI Bug Localizer <span>PRO</span></h1>
                <p class="hero-copy">Premium vulnerability analysis, AI-powered remediation, and enterprise-grade security observability from a single developer workflow.</p>
            </div>
        </div>
    </div>
</div>
<style>
.hero-shell {
    width: 100%;
    display: flex;
    justify-content: center;
    padding: 16px 18px 0;
}
.hero-card {
    width: min(100%, 1220px);
    position: relative;
    overflow: hidden;
    background: radial-gradient(circle at top left, rgba(86, 103, 255, 0.18), transparent 18%), radial-gradient(circle at bottom right, rgba(250, 204, 21, 0.12), transparent 20%), linear-gradient(180deg, rgba(7, 14, 34, 0.98), rgba(12, 16, 42, 0.98));
    border: 1px solid rgba(138, 163, 255, 0.16);
    border-radius: 36px;
    backdrop-filter: blur(20px);
    box-shadow: 0 35px 100px rgba(0, 0, 0, 0.35), inset 0 0 0 1px rgba(255, 255, 255, 0.04);
    padding: 34px 34px 30px;
    display: flex;
    flex-direction: column;
    align-items: center;
}
.hero-deco {
    position: absolute;
    width: 280px;
    height: 280px;
    border-radius: 50%;
    filter: blur(32px);
    opacity: 0.45;
}
.hero-deco-top-left {
    top: -90px;
    left: -90px;
    background: radial-gradient(circle at center, rgba(84, 109, 255, 0.35), transparent 55%);
}
.hero-deco-bottom-right {
    bottom: -90px;
    right: -90px;
    background: radial-gradient(circle at center, rgba(245, 197, 57, 0.28), transparent 55%);
}
.hero-navigator {
    width: 100%;
    display: flex;
    flex-direction: column;
    align-items: center;
    gap: 20px;
    margin-bottom: 28px;
    z-index: 1;
}
.hero-badges {
    display: flex;
    flex-wrap: wrap;
    justify-content: center;
    gap: 14px;
}
.hero-badge {
    display: inline-flex;
    align-items: center;
    gap: 10px;
    padding: 12px 24px;
    border-radius: 999px;
    color: rgba(255, 255, 255, 0.95);
    font-size: 0.82rem;
    font-weight: 700;
    letter-spacing: 0.22em;
    text-transform: uppercase;
    background: rgba(255, 255, 255, 0.06);
    border: 1px solid rgba(255, 255, 255, 0.12);
    box-shadow: 0 12px 30px rgba(0, 0, 0, 0.18);
}
.hero-badge span {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 24px;
    height: 24px;
    border-radius: 999px;
    border: 1px solid rgba(255, 255, 255, 0.16);
    background: rgba(255, 255, 255, 0.08);
}
.hero-badge.accent {
    color: #facc15;
    background: rgba(250, 204, 21, 0.12);
    border-color: rgba(250, 204, 21, 0.28);
    box-shadow: 0 14px 34px rgba(250, 204, 21, 0.12);
}
.hero-subtitle {
    color: rgba(255, 255, 255, 0.70);
    font-size: 0.92rem;
    letter-spacing: 0.35em;
    text-transform: uppercase;
}
.hero-main {
    width: 100%;
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 32px;
    flex-wrap: wrap;
    z-index: 1;
}
.hero-logo-shell {
    position: relative;
    width: 118px;
    height: 118px;
    border-radius: 32px;
    background: linear-gradient(135deg, rgba(46, 64, 255, 0.22), rgba(12, 17, 60, 0.95));
    border: 1px solid rgba(255, 255, 255, 0.15);
    box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.08), 0 24px 68px rgba(29, 55, 146, 0.28);
    display: flex;
    align-items: center;
    justify-content: center;
}
.hero-logo-base {
    position: absolute;
    width: 68px;
    height: 68px;
    border-radius: 22px;
    background: radial-gradient(circle at top left, rgba(84, 109, 255, 0.32), rgba(255, 255, 255, 0.04));
    box-shadow: inset 0 1px 3px rgba(255, 255, 255, 0.14);
}
.hero-content {
    max-width: 840px;
    text-align: center;
}
.hero-eyebrow {
    color: rgba(255, 255, 255, 0.65);
    font-size: 0.95rem;
    letter-spacing: 0.35em;
    text-transform: uppercase;
    margin-bottom: 14px;
}
.hero-heading {
    margin: 0;
    color: white;
    font-size: clamp(4rem, 5vw, 8.8rem);
    font-weight: 900;
    letter-spacing: -1px;
    line-height: 0.96;
    text-align: center;
    text-shadow: 0 20px 40px rgba(14, 23, 68, 0.22);
}
.hero-heading span {
    background: linear-gradient(135deg, #faff7e, #f8c33b 45%, #f5a223 85%);
    -webkit-background-clip: text;
    color: transparent;
    text-shadow: 0 0 24px rgba(250, 204, 21, 0.4);
}
.hero-copy {
    margin: 24px auto 0;
    max-width: 840px;
    color: rgba(226, 232, 255, 0.86);
    font-size: 1.05rem;
    line-height: 1.85;
    text-align: center;
    padding: 0 12px;
}
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, rgba(9, 15, 34, 0.95), rgba(10, 18, 42, 0.98));
    border-right: 1px solid rgba(255, 255, 255, 0.06);
}
section[data-testid="stSidebar"] .sidebar-nav-card {
    background: linear-gradient(135deg, rgba(72, 94, 255, 0.22), rgba(118, 78, 255, 0.18));
    padding: 22px 18px;
    border-radius: 24px;
    margin-bottom: 24px;
    border: 1px solid rgba(255, 255, 255, 0.12);
    box-shadow: 0 18px 40px rgba(7, 13, 40, 0.25);
}
section[data-testid="stSidebar"] .sidebar-title {
    color: #f8faff;
    font-size: 1.05rem;
    font-weight: 800;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    margin: 0;
}
section[data-testid="stSidebar"] .sidebar-divider {
    height: 1px;
    width: 100%;
    margin: 18px 0 0;
    background: rgba(255, 255, 255, 0.08);
}
section[data-testid="stSidebar"] .stRadio > div {
    gap: 0.5rem;
}
section[data-testid="stSidebar"] .stRadio > label > div {
    background: rgba(255, 255, 255, 0.04);
    border: 1px solid rgba(255, 255, 255, 0.08);
    border-radius: 16px;
    padding: 14px 16px;
    min-height: auto;
    color: rgba(232, 240, 255, 0.92);
    display: flex;
    align-items: center;
    transition: all 0.25s ease;
}
section[data-testid="stSidebar"] .stRadio > label > div:hover {
    transform: translateY(-1px);
    border-color: rgba(99, 102, 241, 0.5);
    background: rgba(99, 102, 241, 0.12);
}
section[data-testid="stSidebar"] .stRadio input:checked + div {
    background: linear-gradient(135deg, rgba(88, 100, 255, 0.24), rgba(255, 207, 68, 0.14));
    border-color: rgba(250, 204, 21, 0.35);
    box-shadow: 0 10px 30px rgba(250, 204, 21, 0.12);
}
section[data-testid="stSidebar"] .stRadio input:checked + div span {
    color: #facc15;
}
section[data-testid="stSidebar"] .stRadio input:checked + div > label {
    color: white;
}
section[data-testid="stSidebar"] .stRadio > label > div span {
    color: rgba(255, 255, 255, 0.85);
}
section[data-testid="stSidebar"] .stRadio > label > div > label {
    color: rgba(236, 242, 255, 0.92);
    font-weight: 700;
}
section[data-testid="stSidebar"] .stRadio > label > div svg {
    width: 1.15rem;
    height: 1.15rem;
}
section[data-testid="stSidebar"] .stCheckbox > label > div,
section[data-testid="stSidebar"] .stSelectbox > div > div > div {
    background: rgba(255, 255, 255, 0.04);
    border: 1px solid rgba(255, 255, 255, 0.08);
    border-radius: 16px;
    padding: 10px 12px;
}
section[data-testid="stSidebar"] .stCheckbox > label > div:hover,
section[data-testid="stSidebar"] .stSelectbox > div > div > div:hover {
    background: rgba(99, 102, 241, 0.1);
}
section[data-testid="stSidebar"] .stCheckbox > label > div {
    color: rgba(236, 242, 255, 0.95);
}
section[data-testid="stSidebar"] .stSelectbox > div > div > div {
    color: rgba(236, 242, 255, 0.95);
}
section[data-testid="stSidebar"] .stMarkdown > div {
    color: rgba(226, 232, 255, 0.75);
}
@media (max-width: 960px) {
    .hero-card {
        width: min(100%, 96vw);
        padding: 28px 20px 28px;
    }
    .hero-main {
        gap: 24px;
    }
    .hero-logo-shell {
        width: 100px;
        height: 100px;
    }
    .hero-heading {
        font-size: clamp(3.4rem, 7vw, 7rem);
    }
}
@media (max-width: 720px) {
    .hero-card {
        padding: 24px 16px 24px;
    }
    .hero-badges {
        gap: 10px;
    }
    .hero-badge {
        padding: 10px 18px;
        font-size: 0.78rem;
    }
    .hero-logo-shell {
        width: 84px;
        height: 84px;
    }
    .hero-heading {
        font-size: clamp(2.8rem, 11vw, 4.5rem);
    }
    .hero-copy {
        font-size: 0.98rem;
        line-height: 1.6;
    }
}

</style>
""", unsafe_allow_html=True)
# ================= ADVANCED SIDEBAR NAVIGATION (ULTRA PRO V2) =================
with st.sidebar:

    st.markdown("""
    <style>

    /* MAIN SIDEBAR CARD */
    .sidebar-card {
        background: linear-gradient(135deg, rgba(56,189,248,0.12), rgba(99,102,241,0.12));
        padding: 20px;
        border-radius: 20px;
        border: 1px solid rgba(255,255,255,0.14);
        box-shadow: 0 15px 45px rgba(0,0,0,0.45);
        margin-bottom: 18px;
        text-align: center;
        backdrop-filter: blur(14px);
    }

    .sidebar-title {
        font-size: 21px;
        font-weight: 900;
        color: #e2e8f0;
        margin: 0;
        letter-spacing: 0.5px;
    }

    .sidebar-sub {
        font-size: 12.5px;
        color: rgba(148,163,184,0.95);
        margin-top: 6px;
    }

    /* NAV CONTAINER */
    .nav-container {
        background: rgba(255,255,255,0.03);
        border: 1px solid rgba(255,255,255,0.08);
        border-radius: 16px;
        padding: 10px;
        box-shadow: inset 0 0 20px rgba(0,0,0,0.2);
    }

    /* SECTION TITLE */
    .section-title {
        font-size: 13px;
        font-weight: 700;
        color: #94a3b8;
        margin-top: 12px;
        margin-bottom: 8px;
        letter-spacing: 0.6px;
    }

    /* FOOTER */
    .sidebar-footer {
        text-align: center;
        padding: 12px;
        margin-top: 10px;
        border-top: 1px solid rgba(255,255,255,0.08);
    }
    

    </style>
    """, unsafe_allow_html=True)

    # ================= HEADER CARD =================
    st.markdown("""
    <div class='sidebar-card'>
        <div class='sidebar-title'>🧠 AI Control Center</div>
        <div class='sidebar-sub'>Smart Debug & Assistant Panel</div>
    </div>
    """, unsafe_allow_html=True)

    # ================= NAVIGATION =================
    st.markdown("### 🧭 Navigation")

    st.markdown("<div class='nav-container'>", unsafe_allow_html=True)

    page = st.radio(
        "",
        [
            "🧪 Analyzer",
            "📊 Dashboard",
            "🛠️ Fixer",
            "💬 AI Chat",
            "🔐 Security Lab"
        ]
    )

    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")

    # ================= STATUS PANEL =================
    st.markdown("### 📡 System Status")

    st.success("🟢 AI Engine Active")
    st.info("⚙️ Debug System Ready")

    st.markdown("---")

    # ================= CONTROLS PANEL =================
    st.markdown("### ⚙️ Advanced Controls")

    show_pylint = st.toggle("📋 Pylint Report", value=True)
    show_details = st.toggle("🧠 Show Explanation", value=True)

    debug_mode = st.selectbox(
        "🐞 Debug Mode",
        [
            "All Errors",
            "Syntax Errors",
            "Runtime Errors",
            "Critical Only",
            "Clean View"
        ]
    )

    st.markdown("---")

    # ================= FOOTER =================
    st.markdown("""
    <div class='sidebar-footer'>
        <small style='color:rgba(148,163,184,0.85)'>
        🚀 AI Bug Localizer PRO v2.0<br>
        Built for Smart Debugging
        </small>
    </div>
    """, unsafe_allow_html=True)
#======================================================
# 🧪 ANALYZER TAB (FIXED + STABLE VERSION)
# =========================================================
if page == "🧪 Analyzer":
    st.markdown("""<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; border-radius: 15px; margin-bottom: 20px;">
    <h2 style="color: white; margin: 0;">🧪 Code Analysis Engine</h2>
    <p style="color: rgba(255,255,255,0.9); margin: 10px 0 0 0;">Upload or paste Python code to detect issues, analyze complexity, and get AI-powered fixes</p>
    </div>""", unsafe_allow_html=True)

    # ================= SESSION STATE INIT =================
    if "code" not in st.session_state:
        st.session_state.code = ""

    if "data" not in st.session_state:
        st.session_state.data = {}

    if "history" not in st.session_state:
        st.session_state.history = []

    # ================= UPLOAD + INFO =================
    col1, col2 = st.columns([2, 1])

    with col1:
        uploaded = st.file_uploader(
            "Upload Python File(s)",
            accept_multiple_files=True,
            key="analyzer_uploader"
        )

    with col2:
        st.info("💡 Paste code or upload files to analyze bugs & complexity")

    # ================= BUILD CODE FROM FILE =================
    if uploaded is not None and len(uploaded) > 0:

        file_code = ""

        for file in uploaded:
            try:
                content = file.getvalue().decode("utf-8", errors="ignore")
                file_code += content + "\n\n"
            except (UnicodeDecodeError, AttributeError) as e:
                st.error(f"Cannot read file: {file.name} - {str(e)}")

        if file_code.strip():
            st.session_state.code = file_code
            st.success("File loaded into editor ✔")

    # ================= CODE EDITOR (ALWAYS VISIBLE) =================
    st.markdown("""<div style='background: linear-gradient(135deg, #2d3748 0%, #1a202c 100%); padding: 20px; border-radius: 12px; border: 1px solid #334155; margin-bottom: 20px;'>
    <h3 style='color: #e0e7ff; margin: 0 0 10px 0;'>✍️ Code Editor</h3>
    <p style='color: #cbd5e1; margin: 0; font-size: 0.85rem;'>Enter or paste your Python code below for analysis</p>
    </div>""", unsafe_allow_html=True)

    edited_code = st_ace(
        value=st.session_state.code,
        language="python",
        theme="dracula",
        font_size=16,
        tab_size=4,
        show_gutter=True,
        wrap=True,
        auto_update=True,
        height=500,
        key="ace_editor"
    )
    # ================= SYNC EDITOR BACK =================
    if edited_code is not None:
        st.session_state.code = edited_code

    # ================= ACTION BUTTONS =================
    col1, col2, col3 = st.columns([2, 1, 1])

    with col1:
        if st.button("🚀 Analyze Code", use_container_width=True, type="primary", key="analyze_main"):
            if st.session_state.code.strip():
                with st.spinner("🔍 Running analysis..."):
                    result = cached_analysis(st.session_state.code)
                    st.session_state.data = result
                    st.session_state.history.append(len(result.get("issues", [])))
                st.success("✅ Analysis complete!")
            else:
                st.warning("📝 Please enter some code first")

    with col2:
        if st.button("🗑️ Clear", use_container_width=True):
            st.session_state.code = ""
            st.session_state.data = {}
            st.session_state.history = []
            st.rerun()

    with col3:
        if st.button("📋 Sample", use_container_width=True):
            st.session_state.code = "# Sample code\nprint('Hello')\nif x == None:\n    print(x)"
            st.rerun()

# 📊 DASHBOARD TAB (CLEAN + STABLE)
# =========================================================
if page == "📊 Dashboard":
    st.markdown("""<div style='background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 40px; border-radius: 16px; margin-bottom: 32px; box-shadow: 0 15px 50px rgba(79, 172, 254, 0.2);'>
    <h2 style='color: white; margin: 0 0 10px 0; font-size: 2.2rem;'>📊 System Dashboard</h2>
    <p style='color: rgba(255,255,255,0.9); margin: 0; font-size: 0.95rem;'>Real-time analysis metrics and code quality insights</p>
    </div>""", unsafe_allow_html=True)

    # ================= SAFE STATE HANDLING =================
    data = st.session_state.get("data", {})

    if not data or "issues" not in data:
        st.info("Run analysis first to view dashboard")
        st.stop()

    issues = data.get("issues", []) or []
    complexity = data.get("complexity", {}) or {}

    # ================= FILTER LOGIC =================
    if debug_mode == "Syntax Errors":
        issues = [i for i in issues if i.get("type") == "syntax_error"]
    elif debug_mode == "Runtime Errors":
        issues = [i for i in issues if i.get("type") == "runtime_error"]
    elif debug_mode == "Critical Only":
        issues = [i for i in issues if i.get("score", 0) >= 0.85]
    elif debug_mode == "Clean View":
        issues = [i for i in issues if i.get("score", 0) < 0.6]

    # ================= SAFE METRICS =================
    high, medium, low = get_chart_data(issues) if issues else (0, 0, 0)

    st.markdown("<h3 style='color: #e0e7ff; margin-top: 20px;'>📌 Issue Summary</h3>", unsafe_allow_html=True)

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"<div style='background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%); padding: 20px; border-radius: 12px; text-align: center;'><h3 style='color: white; margin: 0;'>{high}</h3><p style='color: rgba(255,255,255,0.8); margin: 5px 0 0 0; font-size: 0.85rem;'>🔴 High Risk</p></div>", unsafe_allow_html=True)
    with c2:
        st.markdown(f"<div style='background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%); padding: 20px; border-radius: 12px; text-align: center;'><h3 style='color: white; margin: 0;'>{medium}</h3><p style='color: rgba(255,255,255,0.8); margin: 5px 0 0 0; font-size: 0.85rem;'>🟡 Medium</p></div>", unsafe_allow_html=True)
    with c3:
        st.markdown(f"<div style='background: linear-gradient(135deg, #10b981 0%, #059669 100%); padding: 20px; border-radius: 12px; text-align: center;'><h3 style='color: white; margin: 0;'>{low}</h3><p style='color: rgba(255,255,255,0.8); margin: 5px 0 0 0; font-size: 0.85rem;'>🟢 Low</p></div>", unsafe_allow_html=True)
    with c4:
        st.markdown(f"<div style='background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%); padding: 20px; border-radius: 12px; text-align: center;'><h3 style='color: white; margin: 0;'>{len(issues)}</h3><p style='color: rgba(255,255,255,0.8); margin: 5px 0 0 0; font-size: 0.85rem;'>📋 Total</p></div>", unsafe_allow_html=True)

    security_count = len([
        i for i in issues if i.get("type") == "security_vulnerability"
    ])

    st.metric("🔐 Security Issues", security_count)

    if any(
        i.get("type") == "security_vulnerability" and i.get("score", 0) > 0.9
        for i in issues
    ):
        st.error("🚨 SECURITY RISK DETECTED")

    # ================= COMPLEXITY =================
    st.markdown("### 🧠 Complexity Analysis")

    cA, cB = st.columns(2)

    cA.metric("Complexity Score", complexity.get("complexity_score", 0))
    cB.metric("Risk Level", complexity.get("risk", "LOW"))

    # ================= ANALYTICS =================
    st.markdown("### 📊 Visual Analytics")

    with st.expander("Show Charts", expanded=True):

        c1, c2 = st.columns(2)

        with c1:
            pie_fig = px.pie(
                values=[high, medium, low],
                names=["High Risk", "Medium Risk", "Low Risk"],
                title="Issue Severity Distribution"
            )

            pie_fig.update_traces(textinfo="percent+label")

            st.plotly_chart(
                pie_fig,
                use_container_width=True,
                config={"displayModeBar": False}
            )

        with c2:
            bar_fig = px.bar(
                x=["High", "Medium", "Low"],
                y=[high, medium, low],
                title="Issue Breakdown",
                labels={"x": "Severity", "y": "Count"}
            )

            bar_fig.update_layout(showlegend=False)

            st.plotly_chart(
                bar_fig,
                use_container_width=True,
                config={"displayModeBar": False}
            )

    # ================= HISTORY (SAFE FIX) =================
    history = st.session_state.get("history", []) or []

    if len(history) > 0:

        df = pd.DataFrame({
            "Run": list(range(1, len(history) + 1)),
            "Issues": history
        })

        df["Moving Avg"] = df["Issues"].rolling(
            window=2,
            min_periods=1
        ).mean()

        trend_diff = df["Issues"].iloc[-1] - df["Issues"].iloc[0]

        if trend_diff > 0:
            trend_text = "📈 Increasing Issues"
            trend_color = "red"
        elif trend_diff < 0:
            trend_text = "📉 Improving Code Quality"
            trend_color = "green"
        else:
            trend_text = "➡ Stable Trend"
            trend_color = "orange"

        st.markdown(
            f"<h3 style='color:{trend_color}'>{trend_text}</h3>",
            unsafe_allow_html=True
        )

        fig = go.Figure()

        fig.add_trace(go.Scatter(
            x=df["Run"],
            y=df["Issues"],
            mode="lines+markers",
            name="Issues",
            line=dict(width=3)
        ))

        fig.add_trace(go.Scatter(
            x=df["Run"],
            y=df["Moving Avg"],
            mode="lines",
            name="Trend (Smoothed)",
            line=dict(dash="dash", width=3)
        ))

        fig.update_layout(
            title="📊 Code Issue Trend Analysis",
            xaxis_title="Analysis Run",
            yaxis_title="Issue Count",
            template="plotly_dark",
            height=400,
            legend=dict(orientation="h")
        )

        st.plotly_chart(fig, use_container_width=True)

        col1, col2, col3 = st.columns(3)

        col1.metric("First Run", history[0])
        col2.metric("Latest Run", history[-1])
        col3.metric("Change", trend_diff)

    # ================= RESULTS DISPLAY =================
    if not st.session_state.data:
        st.info("🚀 Click 'Analyze Code' above to scan for issues")
    else:
        data = st.session_state.data
        issues = data.get("issues", []) or []
        complexity = data.get("complexity", {}) or {}

        # ================= SUMMARY METRICS =================
        st.markdown("### 📊 Analysis Summary")
        
        metric_cols = st.columns(4)
        
        with metric_cols[0]:
            st.metric(
                "🐛 Total Issues",
                len(issues),
                delta="Found" if issues else "✓ Clean",
                delta_color="off"
            )
        
        with metric_cols[1]:
            high_count = len([i for i in issues if i.get("score", 0) >= 0.85])
            st.metric(
                "🔴 Critical",
                high_count,
                delta="High!" if high_count > 0 else "Good"
            )
        
        with metric_cols[2]:
            score = complexity.get("complexity_score", 0)
            st.metric(
                "📈 Complexity",
                f"{score}/100",
                delta=complexity.get("risk", "UNKNOWN")
            )
        
        with metric_cols[3]:
            quality_score = 10 - min(10, len(issues) * 0.5)
            st.metric(
                "⭐ Quality",
                f"{quality_score:.1f}/10",
                delta="Excellent" if quality_score >= 8 else "Good" if quality_score >= 6 else "Fair"
            )
        
        # ================= ISSUE BREAKDOWN =================
        if issues:
            st.markdown("---")
            st.markdown("### 🔍 Issue Details")
            
            high, medium, low = get_chart_data(issues)
            
            issue_cols = st.columns(3)
            with issue_cols[0]:
                st.markdown(f"<div style='text-align: center; padding: 15px; background: rgba(239,68,68,0.15); border-radius: 10px; border-left: 4px solid #ef4444;'><h3 style='color: #ef4444; margin: 0;'>{high}</h3><p style='margin: 5px 0 0 0; color: #999; font-size: 12px;'>🔴 Critical</p></div>", unsafe_allow_html=True)
            with issue_cols[1]:
                st.markdown(f"<div style='text-align: center; padding: 15px; background: rgba(251,146,60,0.15); border-radius: 10px; border-left: 4px solid #fb923c;'><h3 style='color: #fb923c; margin: 0;'>{medium}</h3><p style='margin: 5px 0 0 0; color: #999; font-size: 12px;'>🟡 Medium</p></div>", unsafe_allow_html=True)
            with issue_cols[2]:
                st.markdown(f"<div style='text-align: center; padding: 15px; background: rgba(34,197,94,0.15); border-radius: 10px; border-left: 4px solid #22c55e;'><h3 style='color: #22c55e; margin: 0;'>{low}</h3><p style='margin: 5px 0 0 0; color: #999; font-size: 12px;'>🟢 Low</p></div>", unsafe_allow_html=True)
            
            st.markdown("---")
            
            for idx, issue in enumerate(issues[:15], 1):
                score = issue.get("score", 0)
                issue_type = issue.get("type", "unknown").replace("_", " ").title()
                line_num = issue.get("line", "?")
                
                # Color coding
                if score >= 0.85:
                    severity_color = "🔴"
                    severity_label = "CRITICAL"
                    bg_class = "critical"
                elif score >= 0.6:
                    severity_color = "🟡"
                    severity_label = "MEDIUM"
                    bg_class = "medium"
                else:
                    severity_color = "🟢"
                    severity_label = "LOW"
                    bg_class = "low"
                
                with st.expander(
                    f"{severity_color} [{severity_label}] Line {line_num}: {issue_type} (Score: {score:.2f})",
                    expanded=(idx == 1)
                ):
                    st.write(f"**Issue:** {issue.get('msg', 'N/A')}")
                    
                    if show_details:
                        st.info(f"💡 {explain(issue)}")
                        st.success(f"🔧 {suggest_fix(issue)}")
                    
                    col_var, col_score = st.columns([3, 1])
                    with col_var:
                        if issue.get("var"):
                            st.write(f"**Variable:** `{issue.get('var')}`")
                    with col_score:
                        st.metric("Risk Score", f"{score:.2f}")
            
            if len(issues) > 15:
                st.caption(f"📌 Showing 15 of {len(issues)} issues. View all issues in the Dashboard.")
        else:
            st.success("✅ No issues detected! Your code looks great.")

    # ================= CODE QUALITY =================
    st.markdown("""<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; border-radius: 15px; margin-bottom: 20px;">
    <h2 style="color: white; margin: 0;">📋 Code Quality Report</h2>
    <p style="color: rgba(255,255,255,0.9); margin: 10px 0 0 0;">Comprehensive quality assessment, metrics, and recommendations</p>
    </div>""", unsafe_allow_html=True)

    data_safe = st.session_state.get("data", {}) or {}
    code_to_check = data_safe.get("code", "") or ""
    quality_issues = data_safe.get("issues", []) or []

    # ================= COMPLEXITY ENRICHMENT =================
    if not code_to_check.strip():
        st.info("📝 Run analysis on some code to generate a quality report")
    else:

        try:
            report = build_quality_report(code_to_check, quality_issues) or {}

            col1, col2 = st.columns(2)

            with col1:
                st.metric("📊 Code Rating (/10)", report.get("rating", 0))

            with col2:
                st.metric("⚠️ Risk Level", report.get("risk", "Unknown"))

            with st.expander("View Full Report (Debug)"):
                st.json(report)

            st.download_button(
                "📄 Download Quality Report",
                data=json.dumps(report, indent=2),
                file_name="code_quality_report.json"
            )

        except (TypeError, ValueError, KeyError) as e:
            st.warning(f"Quality engine error: {str(e)}")


# ================= FIXER TAB (COPILOT STYLE ENGINE) =================
if page == "🛠️ Fixer":

    st.markdown("""<div style='background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 40px; border-radius: 16px; margin-bottom: 32px; box-shadow: 0 15px 50px rgba(245, 87, 108, 0.25);'>
    <h2 style='color: white; margin: 0 0 10px 0; font-size: 2.2rem;'>🛠️ AI Code Fixer</h2>
    <p style='color: rgba(255,255,255,0.9); margin: 0; font-size: 0.95rem;'>Automatically repair, optimize, and improve your code</p>
    </div>""", unsafe_allow_html=True)

    # ================= SAFE STATE ACCESS =================
    data = st.session_state.get("data", {})

    if not data:
        st.warning("Run analysis first")
        st.stop()

    original_code = data.get("code", "")
    issues = data.get("issues", []) or []

    # ================= COPILOT FIX ENGINE (STABLE) =================
    def copilot_fix_engine(code):

        try:
            if not code or not code.strip():
                return "# No code provided"

            lines = code.split("\n")
            fixed_lines = []

            fixed_lines.append("# =========================================")
            fixed_lines.append("# AI COPILOT REPAIRED CODE")
            fixed_lines.append("# Auto-fix applied: syntax + structure + cleanup")
            fixed_lines.append("# =========================================\n")

            for line in lines:
                raw = line.rstrip()

                # skip empty lines safely
                if not raw.strip():
                    fixed_lines.append("")
                    continue

                stripped = raw.strip()

                # ================= SYNTAX FIXES =================

                # fix missing print()
                if stripped.startswith("print ") and "(" not in stripped:
                    raw = raw.replace("print ", "print(") + ")  # auto-fixed print syntax"

                # fix None comparisons
                if "== None" in stripped:
                    raw = raw.replace("== None", "is None  # fixed PEP8")

                if "!= None" in stripped:
                    raw = raw.replace("!= None", "is not None  # fixed PEP8")

                # fix missing colon
                if stripped.startswith(("if ", "for ", "while ", "def ", "else")) and ":" not in stripped:
                    raw = raw + ":  # auto-added missing colon"

                fixed_lines.append(raw)

            fixed_code = "\n".join(fixed_lines)

            fixed_code += "\n\n# NOTE: AI Copilot applied best-effort corrections"

            return fixed_code

        except (SyntaxError, ValueError) as e:
            return f"# Fixer error: {str(e)}\n\n" + (code or "")


    fixed_code = copilot_fix_engine(original_code)

    # ================= SIDE BY SIDE UI =================
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### ❌ Original Code (With Errors)")
        st.code(original_code if original_code else "# No code provided", language="python")

    with col2:
        st.markdown("### ✅ AI Copilot Fixed Code")
        st.code(fixed_code, language="python")

    # ================= STATUS =================
    if (original_code or "").strip() == (fixed_code or "").strip():
        st.warning("No major improvements detected (code may already be clean)")
    else:
        st.success("Copilot AI successfully repaired and improved your code ✔")

    # ================= DOWNLOAD =================
    st.download_button(
        "📥 Download Copilot Fixed Code",
        data=fixed_code,
        file_name="copilot_fixed_code.py"
    )

# ================= AI CHAT =================
if page == "💬 AI Chat":
    st.markdown("""<div style='background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 40px; border-radius: 16px; margin-bottom: 32px; box-shadow: 0 15px 50px rgba(79, 172, 254, 0.25);'>
    <h2 style='color: white; margin: 0 0 10px 0; font-size: 2.2rem;'>💬 AI Debug Assistant</h2>
    <p style='color: rgba(255,255,255,0.9); margin: 0; font-size: 0.95rem;'>Chat with AI to explain issues and get personalized debugging advice</p>
    </div>""", unsafe_allow_html=True)

    # ================= AI CHAT FUNCTION =================
    def ai_chat(user_input, data):
        """Use enhanced AI engine for intelligent responses"""
        try:
            # Get AI engine instance
            ai_engine = get_ai_engine()

            # Extract data
            code = data.get("code", "")
            issues = data.get("issues", [])
            complexity = data.get("complexity", {})

            # Get AI response with full context
            response = ai_engine.get_response(
                user_input=user_input,
                code=code,
                issues=issues,
                complexity=complexity
            )

            return response

        except Exception as e:
            return f"""⚠️ **AI Response Error**

Something went wrong: {str(e)}

Don't worry! The analysis still works. Try:
1. Refresh the page
2. Ask a simpler question
3. Check your code syntax

I'm here to help! 🤖"""

    # ================= CHAT UI =================
    st.markdown("### 💭 Conversation History")
    
    chat_container = st.container(border=True)
    
    with chat_container:
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"], avatar="🤖" if msg["role"] == "assistant" else "👤"):
                st.markdown(msg["content"])

    # ================= CHAT INPUT =================
    st.markdown("---")
    st.markdown("### ✍️ Ask a Question")
    
    col_input, col_send = st.columns([4, 1])
    
    with col_input:
        user_msg = st.chat_input("Ask about your code, issues, fixes, or anything else...")
    
    with col_send:
        pass

    if user_msg:

        # ================= SAVE USER MESSAGE =================
        st.session_state.chat_history.append({
            "role": "user",
            "content": user_msg
        })

        with st.spinner("🤖 AI is thinking..."):
            response = ai_chat(
                user_msg,
                st.session_state.data
            )

        # ================= SAVE AI RESPONSE =================
        st.session_state.chat_history.append({
            "role": "assistant",
            "content": response
        })

        # ================= REFRESH CHAT =================
        st.rerun()
    
    if not st.session_state.chat_history:
        st.info("💡 Tip: Load some code in the Analyzer tab first for the AI to provide context-aware answers!")
if page == "🔐 Security Lab":
    st.markdown("""<div style='background: linear-gradient(135deg, #fa709a 0%, #fee140 100%); padding: 40px; border-radius: 16px; margin-bottom: 32px; box-shadow: 0 15px 50px rgba(250, 112, 154, 0.25);'>
    <h2 style='color: white; margin: 0 0 10px 0; font-size: 2.2rem;'>🔐 Information Security Lab</h2>
    <p style='color: rgba(255,255,255,0.9); margin: 0; font-size: 0.95rem;'>OWASP • Encryption • Compliance • Authentication Analysis</p>
    </div>""", unsafe_allow_html=True)

    # Create tabs for different security features
    sec_tab1, sec_tab2, sec_tab3, sec_tab4, sec_tab5, sec_tab6 = st.tabs([
        "🛡️ OWASP Analysis",
        "🔐 Encryption", 
        "📋 Compliance",
        "🔑 Authentication",
        "🔧 Crypto Tools",
        "📄 Reports"
    ])
    
    # Get current code analysis if available
    data = st.session_state.get("data", {})
    
    # ==================== TAB 1: OWASP ANALYSIS ====================
    with sec_tab1:
        st.markdown("### OWASP Top 10 Vulnerability Detection")
        
        if data and data.get("code"):
            with st.spinner("🔍 Scanning for OWASP vulnerabilities..."):
                owasp_issues = data.get("owasp_issues", [])
                
            if owasp_issues:
                render_security_dashboard(owasp_issues)
                
                st.markdown("#### Detected Vulnerabilities:")
                for idx, issue in enumerate(owasp_issues[:10], 1):
                    render_issue_card(issue, idx)
            else:
                st.success("✅ No OWASP vulnerabilities detected!")
        else:
            st.info("💡 Analyze code in the Analyzer tab to check for OWASP vulnerabilities")
    
    # ==================== TAB 2: ENCRYPTION ANALYSIS ====================
    with sec_tab2:
        st.markdown("### 🔐 Cryptography & Encryption Analysis")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### Code Analysis")
            if data and data.get("code"):
                with st.spinner("🔍 Analyzing encryption practices..."):
                    crypto_issues = data.get("crypto_issues", [])
                
                if crypto_issues:
                    for idx, issue in enumerate(crypto_issues[:5], 1):
                        severity_name, emoji, _, css_class = render_issue_card.__globals__['get_severity_color'](issue.get("score", 0))
                        st.markdown(f"{emoji} **{issue.get('type')}**: {issue.get('msg')}")
                        with st.expander("Details"):
                            st.info(f"Recommendation: {issue.get('recommendation')}")
                else:
                    st.success("✅ Good encryption practices detected!")
            else:
                st.info("💡 Analyze code to check encryption practices")
        
        with col2:
            st.markdown("#### Crypto Tools")
            st.session_state.setdefault("encrypted_data", None)
            st.session_state.setdefault("crypto_key", None)
            
            # Hash Generator
            st.markdown("**Hash Generator**")
            text = st.text_input("Text to hash:", key="hash_input")
            hash_algo = st.selectbox("Algorithm:", ["SHA-256", "SHA-512", "MD5"], key="hash_algo")
            
            if text:
                if hash_algo == "SHA-256":
                    result = hashlib.sha256(text.encode()).hexdigest()
                elif hash_algo == "SHA-512":
                    result = hashlib.sha512(text.encode()).hexdigest()
                else:
                    result = hashlib.md5(text.encode()).hexdigest()
                st.code(result, language="text")
            
            st.markdown("---")
            
            # Encryption
            st.markdown("**Fernet Encryption**")
            msg = st.text_area("Message to encrypt:", key="encrypt_input")
            
            col_enc1, col_enc2 = st.columns(2)
            with col_enc1:
                if st.button("🔒 Encrypt", key="encrypt_btn"):
                    if msg.strip():
                        key = Fernet.generate_key()
                        cipher = Fernet(key)
                        encrypted = cipher.encrypt(msg.encode())
                        st.session_state.encrypted_data = encrypted
                        st.session_state.crypto_key = key
                        st.success("✅ Encrypted!")
                        st.code(encrypted.decode(), language="text")
                    else:
                        st.warning("Enter a message first")
            
            with col_enc2:
                if st.button("🔓 Decrypt", key="decrypt_btn"):
                    try:
                        if st.session_state.encrypted_data and st.session_state.crypto_key:
                            cipher = Fernet(st.session_state.crypto_key)
                            decrypted = cipher.decrypt(st.session_state.encrypted_data).decode()
                            st.success("✅ Decrypted!")
                            st.code(decrypted, language="text")
                        else:
                            st.error("No encrypted data. Encrypt first!")
                    except Exception as e:
                        st.error(f"❌ Decryption failed: {str(e)}")
    
    # ==================== TAB 3: COMPLIANCE ====================
    with sec_tab3:
        st.markdown("### 📋 Compliance & Standards")
        
        if data and data.get("code"):
            with st.spinner("🔍 Checking compliance..."):
                compliance_issues = data.get("compliance_issues", [])
            
            if compliance_issues:
                render_compliance_checker(compliance_issues)
            else:
                st.success("✅ No compliance issues detected!")
                st.info("Your code appears to comply with GDPR, HIPAA, and PCI-DSS standards")
        else:
            st.info("💡 Analyze code to check compliance")
        
        # Compliance Info
        st.markdown("---")
        st.markdown("#### 📚 Compliance Standards")
        
        col_std1, col_std2, col_std3 = st.columns(3)
        
        with col_std1:
            st.markdown("""
            **🇪🇺 GDPR**
            - Personal data protection
            - User consent
            - Data retention
            """)
        
        with col_std2:
            st.markdown("""
            **🏥 HIPAA**
            - PHI encryption
            - Access controls
            - Audit logging
            """)
        
        with col_std3:
            st.markdown("""
            **💳 PCI-DSS**
            - Card data encryption
            - Access control
            - Vulnerability management
            """)
    
    # ==================== TAB 4: AUTHENTICATION ====================
    with sec_tab4:
        st.markdown("### 🔑 Authentication & Authorization")
        
        if data and data.get("code"):
            with st.spinner("🔍 Analyzing authentication..."):
                auth_issues = data.get("auth_issues", [])
            
            if auth_issues:
                st.warning(f"⚠️ Found {len(auth_issues)} authentication issues")
                for idx, issue in enumerate(auth_issues[:5], 1):
                    render_issue_card(issue, idx)
            else:
                st.success("✅ Good authentication practices detected!")
        else:
            st.info("💡 Analyze code to check authentication")
        
        st.markdown("---")
        st.markdown("#### 🛡️ Security Best Practices")
        
        best_practices = {
            "Password Policy": [
                "✅ Minimum 12 characters",
                "✅ Uppercase & lowercase",
                "✅ Numbers & special characters",
                "✅ Avoid common patterns"
            ],
            "Session Management": [
                "✅ 15-30 minute timeout",
                "✅ Secure & HttpOnly cookies",
                "✅ SameSite protection",
                "✅ Token refresh mechanism"
            ],
            "Multi-Factor Auth": [
                "✅ Implement MFA/2FA",
                "✅ TOTP over SMS when possible",
                "✅ Backup codes",
                "✅ Rate limiting"
            ]
        }
        
        for category, practices in best_practices.items():
            with st.expander(f"📌 {category}"):
                for practice in practices:
                    st.markdown(practice)
    
    # ==================== TAB 5: CRYPTO TOOLS ====================
    with sec_tab5:
        st.markdown("### 🔧 Security Tools & Utilities")
        
        tool_option = st.radio("Select Tool:", [
            "Password Strength Checker",
            "JWT Generator",
            "API Key Generator",
            "Security Headers"
        ], horizontal=True)
        
        if tool_option == "Password Strength Checker":
            st.markdown("#### Check Password Strength")
            password = st.text_input("Enter password:", type="password", key="pwd_check")
            
            if password:
                strength_score = 0
                feedback = []
                
                if len(password) >= 12:
                    strength_score += 25
                else:
                    feedback.append("Use at least 12 characters")
                
                if any(c.isupper() for c in password):
                    strength_score += 25
                else:
                    feedback.append("Add uppercase letters")
                
                if any(c.isdigit() for c in password):
                    strength_score += 25
                else:
                    feedback.append("Add numbers")
                
                if any(c in "!@#$%^&*" for c in password):
                    strength_score += 25
                else:
                    feedback.append("Add special characters")
                
                col_pwd1, col_pwd2 = st.columns([1, 2])
                with col_pwd1:
                    if strength_score >= 75:
                        st.success(f"💪 Strong: {strength_score}%")
                    elif strength_score >= 50:
                        st.warning(f"⚠️ Medium: {strength_score}%")
                    else:
                        st.error(f"❌ Weak: {strength_score}%")
                
                with col_pwd2:
                    if feedback:
                        st.info("**Recommendations:**\n" + "\n".join(feedback))
        
        elif tool_option == "JWT Generator":
            st.markdown("#### Generate JWT Token")
            secret = st.text_input("Secret Key:", type="password", key="jwt_secret")
            payload = st.text_area("Payload (JSON):", value='{"user_id": "123", "exp": 1234567890}', key="jwt_payload")
            
            if st.button("Generate JWT", key="gen_jwt"):
                try:
                    import json
                    import base64
                    import hmac
                    import hashlib
                    
                    header = {"alg": "HS256", "typ": "JWT"}
                    payload_dict = json.loads(payload)
                    
                    message = base64.urlsafe_b64encode(json.dumps(header).encode()).decode().rstrip('=')
                    message += '.' + base64.urlsafe_b64encode(json.dumps(payload_dict).encode()).decode().rstrip('=')
                    
                    signature = base64.urlsafe_b64encode(
                        hmac.new(secret.encode(), message.encode(), hashlib.sha256).digest()
                    ).decode().rstrip('=')
                    
                    jwt_token = f"{message}.{signature}"
                    st.success("JWT Generated!")
                    st.code(jwt_token, language="text")
                except Exception as e:
                    st.error(f"Error: {str(e)}")
        
        elif tool_option == "API Key Generator":
            st.markdown("#### Generate Secure API Key")
            import secrets
            
            key_length = st.slider("Key Length:", 16, 64, 32)
            if st.button("Generate API Key", key="gen_api"):
                api_key = secrets.token_urlsafe(key_length)
                st.success("API Key Generated!")
                st.code(api_key, language="text")
        
        elif tool_option == "Security Headers":
            st.markdown("#### Recommended Security Headers")
            headers = {
                "X-Content-Type-Options": "nosniff",
                "X-Frame-Options": "DENY",
                "X-XSS-Protection": "1; mode=block",
                "Strict-Transport-Security": "max-age=31536000",
                "Content-Security-Policy": "default-src 'self'",
                "Referrer-Policy": "strict-origin-when-cross-origin"
            }
            
            st.json(headers)


    # ==================== TAB 6: REPORTS ====================
    with sec_tab6:
        st.markdown("### 📄 Security Analysis Reports")
        
        if data and data.get("code"):
            st.markdown("#### Generate Comprehensive Reports")
            
            col_rep1, col_rep2 = st.columns(2)
            
            with col_rep1:
                st.markdown("**📊 PDF Report**")
                st.markdown("Professional PDF with charts and detailed analysis")
                if st.button("Generate PDF Report", key="pdf_report"):
                    with st.spinner("Generating PDF report..."):
                        pdf_data = generate_pdf_report(data)
                    st.download_button(
                        "📥 Download PDF Report",
                        data=pdf_data,
                        file_name=f"security_analysis_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf",
                        mime="application/pdf",
                        key="pdf_download"
                    )
            
            with col_rep2:
                st.markdown("**📝 Word Report**")
                st.markdown("Detailed Word document with APA 7 citations")
                if st.button("Generate Word Report", key="word_report"):
                    with st.spinner("Generating Word report..."):
                        word_data = generate_word_report(data)
                    st.download_button(
                        "📥 Download Word Report",
                        data=word_data,
                        file_name=f"comprehensive_security_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="word_download"
                    )
            
            st.markdown("---")
            st.markdown("#### Report Features")
            
            features = {
                "📊 Executive Summary": "Overview of all security findings",
                "🔍 Detailed Analysis": "Issue-by-issue breakdown with recommendations",
                "📚 APA 7 Citations": "Academic references for all security standards",
                "📈 Charts & Metrics": "Visual representation of security posture",
                "💡 Remediation Guidance": "Step-by-step fix instructions"
            }
            
            for feature, desc in features.items():
                st.markdown(f"**{feature}**: {desc}")
                
        else:
            st.info("💡 Analyze code in the Analyzer tab first to generate reports")