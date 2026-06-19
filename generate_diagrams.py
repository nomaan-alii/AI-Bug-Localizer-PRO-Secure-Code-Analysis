"""
Generate architecture and flow diagrams for AI Bug Localizer Project
Creates a comprehensive Word document with Mermaid diagrams
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import os
from datetime import datetime

def add_heading_with_style(doc, text, level, color=None):
    """Add styled heading to document"""
    heading = doc.add_heading(text, level=level)
    if color:
        heading.runs[0].font.color.rgb = color
    return heading

def add_colored_paragraph(doc, text, size=11, bold=False, color=None, alignment=None):
    """Add styled paragraph"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    return p

def add_table_with_data(doc, headers, data):
    """Add formatted table to document"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def create_mermaid_diagram(diagram_code, output_file):
    """Create Mermaid diagram as SVG"""
    try:
        # Using mermaid.cli if available
        with open('temp_mermaid.md', 'w') as f:
            f.write(f"```mermaid\n{diagram_code}\n```")
        
        # If mmdc is available
        subprocess.run(['mmdc', '-i', 'temp_mermaid.md', '-o', output_file], 
                      capture_output=True, timeout=10)
        if os.path.exists(output_file):
            return output_file
    except:
        pass
    return None

def generate_architecture_diagram():
    """Generate architecture diagram as text description"""
    return """
ARCHITECTURE DIAGRAM - AI BUG LOCALIZER PRO

┌─────────────────────────────────────────────────────────────────┐
│                     STREAMLIT UI LAYER                          │
│  ┌──────────────┐ ┌──────────────┐ ┌──────────────────────────┐ │
│  │ 🧪 Analyzer │ │ 🔐 Security  │ │ 💬 AI Chat Assistant   │ │
│  │   Tab       │ │   Lab Tab    │ │                          │ │
│  └──────────────┘ └──────────────┘ └──────────────────────────┘ │
└────────────────────────────┬────────────────────────────────────┘
                             │
┌────────────────────────────▼────────────────────────────────────┐
│            ANALYSIS ENGINE LAYER                                 │
│  ┌─────────────────────────────────────────────────────────────┐│
│  │ STATIC ANALYSIS           RUNTIME ANALYSIS   COMPLEXITY      ││
│  │ • AST Parser              • Error Detection  • Metrics        ││
│  │ • Pattern Matching        • Execution Flow   • Scoring        ││
│  └─────────────────────────────────────────────────────────────┘│
└────────────────────────────┬────────────────────────────────────┘
                             │
┌────────────────────────────▼────────────────────────────────────┐
│         SECURITY ANALYSIS MODULES LAYER                          │
│  ┌──────────────┐ ┌──────────────┐ ┌──────────────────────────┐ │
│  │ OWASP        │ │ Encryption   │ │ Compliance Checker       │ │
│  │ Top 10       │ │ Analyzer     │ │ (GDPR, HIPAA, PCI-DSS)   │ │
│  └──────────────┘ └──────────────┘ └──────────────────────────┘ │
│  ┌──────────────┐ ┌──────────────┐ ┌──────────────────────────┐ │
│  │ Authentication│ │ API Security │ │ Dependency               │ │
│  │ Analyzer     │ │ Analyzer     │ │ Vulnerability Analyzer   │ │
│  └──────────────┘ └──────────────┘ └──────────────────────────┘ │
│  ┌──────────────┐ ┌──────────────┐ ┌──────────────────────────┐ │
│  │ Config       │ │ Input        │ │ Threat Modeling          │ │
│  │ Analyzer     │ │ Validation   │ │ Analyzer                 │ │
│  └──────────────┘ └──────────────┘ └──────────────────────────┘ │
└────────────────────────────┬────────────────────────────────────┘
                             │
┌────────────────────────────▼────────────────────────────────────┐
│             AI & ML ENGINE LAYER                                │
│  ┌─────────────────────────────────────────────────────────────┐│
│  │ RANKING ENGINE        FIX SUGGESTER        AI CHAT ENGINE    ││
│  │ • ML-based scoring    • Code corrections  • Context-aware    ││
│  │ • Prioritization      • Recommendations   • Explanations     ││
│  └─────────────────────────────────────────────────────────────┘│
└────────────────────────────┬────────────────────────────────────┘
                             │
┌────────────────────────────▼────────────────────────────────────┐
│              CORE PROCESSING LAYER                               │
│  ┌──────────────────────────────────────────────────────────────┐│
│  │ QUALITY ENGINE        REPORT ENGINE       CODE CORRECTOR     ││
│  │ • Metrics calculation • Report generation • Auto-fix engine  ││
│  │ • KPI tracking        • PDF/Export        • PEP8 cleanup     ││
│  └──────────────────────────────────────────────────────────────┘│
└────────────────────────────┬────────────────────────────────────┘
                             │
┌────────────────────────────▼────────────────────────────────────┐
│             OUTPUT & VISUALIZATION LAYER                         │
│  ┌──────────────┐ ┌──────────────┐ ┌──────────────────────────┐ │
│  │ Dashboard    │ │ PDF Reports  │ │ Issue Cards & Charts     │ │
│  │ Components   │ │              │ │ (Plotly Visualization)   │ │
│  └──────────────┘ └──────────────┘ └──────────────────────────┘ │
└────────────────────────────────────────────────────────────────┘
"""

def generate_flow_diagram():
    """Generate main flow diagram"""
    return """
OVERALL WORKFLOW DIAGRAM - AI BUG LOCALIZER PRO

START: User opens application
         │
         ▼
    ┌─────────────┐
    │ Select Tab  │
    │  & Feature  │
    └──────┬──────┘
           │
    ┌──────┴──────────────┬──────────────┐
    │                     │              │
    ▼                     ▼              ▼
 CODE INPUT         SECURITY LAB    AI CHAT
    │                   │              │
    ├─ Paste Code       │              │
    ├─ Upload File      │              │
    └─┬────────────────┘│              │
      │                 │              │
      ▼                 ▼              ▼
   ┌──────────────────────────────────────┐
   │  CODE ANALYSIS PIPELINE              │
   │  1. Static Analysis (AST)            │
   │  2. Runtime Analysis                 │
   │  3. Complexity Calculation           │
   └────────────┬─────────────────────────┘
                │
                ▼
   ┌──────────────────────────────────────┐
   │  SECURITY ANALYSIS PIPELINE          │
   │  • OWASP Top 10 Detection            │
   │  • Encryption Analysis               │
   │  • Compliance Checking               │
   │  • Authentication Review             │
   │  • API Security                      │
   │  • Dependency Vulnerabilities        │
   │  • Configuration Check               │
   │  • Input Validation                  │
   │  • Threat Modeling                   │
   └────────────┬─────────────────────────┘
                │
                ▼
   ┌──────────────────────────────────────┐
   │  ISSUE RANKING & SCORING             │
   │  • ML-based Ranking                  │
   │  • Severity Classification           │
   │  • Priority Assignment               │
   │  • Risk Assessment                   │
   └────────────┬─────────────────────────┘
                │
                ▼
   ┌──────────────────────────────────────┐
   │  FIX GENERATION & EXPLANATION        │
   │  • Auto Fix Suggestions              │
   │  • Code Corrections                  │
   │  • Detailed Explanations             │
   │  • Recommendations                   │
   └────────────┬─────────────────────────┘
                │
                ▼
   ┌──────────────────────────────────────┐
   │  REPORT GENERATION                   │
   │  • Dashboard Visualization           │
   │  • PDF Report Export                 │
   │  • Quality Metrics                   │
   │  • KPI Summary                       │
   └────────────┬─────────────────────────┘
                │
                ▼
   ┌──────────────────────────────────────┐
   │  DISPLAY RESULTS TO USER             │
   │  • Issue Cards                       │
   │  • Charts & Graphs                   │
   │  • Fix Suggestions                   │
   │  • Download Report                   │
   └────────────┬─────────────────────────┘
                │
                ▼
           END OF CYCLE
"""

def generate_security_pipeline():
    """Generate security analysis pipeline diagram"""
    return """
SECURITY ANALYSIS PIPELINE - DETAILED VIEW

INPUT: Code
  │
  ▼
┌──────────────────────────────────────────┐
│ OWASP TOP 10 ANALYZER                    │
│ • SQL Injection Detection                │
│ • Cross-Site Scripting (XSS)            │
│ • Authentication Flaws                   │
│ • Sensitive Data Exposure                │
│ • XML External Entity (XXE)             │
│ • Broken Access Control                  │
│ • Security Misconfiguration             │
│ • Insecure Deserialization              │
│ • Known Vulnerabilities                  │
│ • Insufficient Logging                   │
└──────────┬───────────────────────────────┘
           │
           ▼
┌──────────────────────────────────────────┐
│ ENCRYPTION ANALYZER                      │
│ • Weak Algorithm Detection               │
│ • SSL/TLS Validation                    │
│ • Password Hashing Security              │
│ • Key Management Issues                  │
│ • Cipher Strength Analysis               │
└──────────┬───────────────────────────────┘
           │
           ▼
┌──────────────────────────────────────────┐
│ COMPLIANCE CHECKER                       │
│ • GDPR Requirements                      │
│ • HIPAA Standards                        │
│ • PCI-DSS Compliance                     │
│ • Data Protection Rules                  │
└──────────┬───────────────────────────────┘
           │
           ▼
┌──────────────────────────────────────────┐
│ AUTHENTICATION ANALYZER                  │
│ • Password Strength                      │
│ • Session Management                     │
│ • MFA Implementation                     │
│ • JWT Security                           │
│ • OAuth Configuration                    │
└──────────┬───────────────────────────────┘
           │
           ▼
┌──────────────────────────────────────────┐
│ API SECURITY ANALYZER                    │
│ • Endpoint Security                      │
│ • Rate Limiting                          │
│ • API Authentication                     │
│ • CORS Configuration                     │
└──────────┬───────────────────────────────┘
           │
           ▼
┌──────────────────────────────────────────┐
│ DEPENDENCY VULNERABILITY ANALYZER        │
│ • Known Vulnerabilities                  │
│ • Outdated Packages                      │
│ • Security Patches                       │
│ • Compatibility Issues                   │
└──────────┬───────────────────────────────┘
           │
           ▼
┌──────────────────────────────────────────┐
│ CONFIG SECURITY ANALYZER                 │
│ • Security Settings                      │
│ • Debug Mode Detection                   │
│ • Secrets in Code                        │
│ • Hardcoded Credentials                  │
└──────────┬───────────────────────────────┘
           │
           ▼
┌──────────────────────────────────────────┐
│ INPUT VALIDATION ANALYZER                │
│ • Input Sanitization                     │
│ • Type Checking                          │
│ • Boundary Validation                    │
│ • Format Verification                    │
└──────────┬───────────────────────────────┘
           │
           ▼
┌──────────────────────────────────────────┐
│ THREAT MODELING ANALYZER                 │
│ • Risk Assessment                        │
│ • Attack Vectors                         │
│ • Mitigation Strategies                  │
│ • Impact Analysis                        │
└──────────┬───────────────────────────────┘
           │
           ▼
OUTPUT: Aggregated Security Issues & Recommendations
"""

def create_word_document():
    """Create comprehensive Word document with diagrams"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title Page
    title = doc.add_heading('AI BUG LOCALIZER PRO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    subtitle = doc.add_heading('Architecture & Flow Diagrams', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = add_colored_paragraph(doc, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_paragraph()
    
    # Table of Contents
    add_heading_with_style(doc, "📋 Table of Contents", 1, RGBColor(0, 102, 204))
    toc_items = [
        "1. Project Overview",
        "2. System Architecture Diagram",
        "3. Overall Workflow Diagram",
        "4. Security Analysis Pipeline",
        "5. Component Description",
        "6. Data Flow Architecture",
        "7. Technology Stack"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Project Overview
    add_heading_with_style(doc, "1. Project Overview", 1, RGBColor(0, 102, 204))
    overview_text = """AI Bug Localizer PRO is an enterprise-grade AI-powered Python debugging and security validation 
system. It combines advanced static and runtime analysis with machine learning ranking to identify bugs, security 
vulnerabilities, and compliance issues in Python code."""
    doc.add_paragraph(overview_text)
    
    add_heading_with_style(doc, "Key Capabilities", 2)
    capabilities = [
        "AST-based static code analysis",
        "OWASP Top 10 vulnerability detection",
        "Cryptographic analysis and validation",
        "Compliance checking (GDPR, HIPAA, PCI-DSS)",
        "ML-based issue ranking and prioritization",
        "Automated fix suggestions",
        "Interactive AI debugging assistant",
        "PDF report generation",
        "Real-time security dashboard"
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_page_break()
    
    # Architecture Diagram
    add_heading_with_style(doc, "2. System Architecture Diagram", 1, RGBColor(0, 102, 204))
    doc.add_paragraph("The following diagram shows the layered architecture of AI Bug Localizer PRO:")
    doc.add_paragraph()
    
    arch_diagram = generate_architecture_diagram()
    for line in arch_diagram.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
    
    doc.add_paragraph()
    add_colored_paragraph(doc, "Architecture Layers Explained:", bold=True)
    
    layers = [
        ("Streamlit UI Layer", "Provides interactive user interface for code analysis and security labs"),
        ("Analysis Engine Layer", "Performs static, runtime, and complexity analysis"),
        ("Security Modules Layer", "Houses all security-specific analyzers"),
        ("AI & ML Engine Layer", "Handles issue ranking, fix suggestions, and AI chat"),
        ("Core Processing Layer", "Quality metrics, report generation, and code correction"),
        ("Output & Visualization Layer", "Dashboard components and report export")
    ]
    
    for layer_name, description in layers:
        p = doc.add_paragraph()
        p.add_run(layer_name).bold = True
        p.add_run(f": {description}")
    
    doc.add_page_break()
    
    # Overall Workflow
    add_heading_with_style(doc, "3. Overall Workflow Diagram", 1, RGBColor(0, 102, 204))
    doc.add_paragraph("The following diagram illustrates the main workflow from user input to results:")
    doc.add_paragraph()
    
    flow_diagram = generate_flow_diagram()
    for line in flow_diagram.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Security Pipeline
    add_heading_with_style(doc, "4. Security Analysis Pipeline", 1, RGBColor(0, 102, 204))
    doc.add_paragraph("Detailed view of how code flows through the security analysis pipeline:")
    doc.add_paragraph()
    
    security_pipeline = generate_security_pipeline()
    for line in security_pipeline.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Component Description
    add_heading_with_style(doc, "5. Component Description", 1, RGBColor(0, 102, 204))
    
    components_data = [
        ("Component", "Purpose", "Key Files"),
        ("Static Analysis", "AST-based code pattern matching", "static_analysis.py"),
        ("Runtime Analysis", "Error detection during execution", "runtime_analysis.py"),
        ("OWASP Analyzer", "OWASP Top 10 vulnerability detection", "owasp_analyzer.py"),
        ("Encryption Analyzer", "Cryptographic weakness detection", "encryption_analyzer.py"),
        ("Compliance Checker", "GDPR, HIPAA, PCI-DSS validation", "compliance_checker.py"),
        ("Authentication Analyzer", "Auth mechanism security review", "authentication_analyzer.py"),
        ("API Security Analyzer", "REST API security assessment", "api_security_analyzer.py"),
        ("Dependency Analyzer", "Vulnerable package detection", "dependency_vulnerability_analyzer.py"),
        ("Config Analyzer", "Security configuration review", "secure_configuration_analyzer.py"),
        ("Input Validation", "Input sanitization checking", "input_validation_analyzer.py"),
        ("Threat Modeling", "Risk and attack vector analysis", "threat_modeling_analyzer.py"),
        ("ML Ranking Engine", "Issue prioritization via ML", "ranking.py"),
        ("Fix Suggester", "Automated fix recommendations", "fix_suggester.py"),
        ("AI Chat Engine", "Context-aware assistance", "enhanced_ai_chat.py")
    ]
    
    add_table_with_data(doc, components_data[0], components_data[1:])
    
    doc.add_page_break()
    
    # Data Flow Architecture
    add_heading_with_style(doc, "6. Data Flow Architecture", 1, RGBColor(0, 102, 204))
    
    add_heading_with_style(doc, "Code Input Processing", 2)
    data_flow_input = """
1. User Input → Code Editor (Streamlit ACE)
2. Code Stored → In-memory storage
3. Preprocessing → Syntax validation
4. AST Generation → Abstract Syntax Tree creation
5. Tokenization → Code token extraction
    """
    doc.add_paragraph(data_flow_input)
    
    add_heading_with_style(doc, "Analysis Processing", 2)
    analysis_flow = """
1. Static Analysis → Pattern matching on AST
2. Runtime Analysis → Execute code with error handling
3. Security Modules → Run all security checks in parallel
4. Aggregation → Combine all findings
5. De-duplication → Remove redundant issues
    """
    doc.add_paragraph(analysis_flow)
    
    add_heading_with_style(doc, "Results Processing", 2)
    results_flow = """
1. Issue Ranking → ML model scores and ranks issues
2. Severity Classification → HIGH, MEDIUM, LOW, CRITICAL
3. Fix Generation → Automated fix suggestions
4. Explanation Generation → AI-powered descriptions
5. Report Creation → Dashboard rendering & PDF export
    """
    doc.add_paragraph(results_flow)
    
    doc.add_page_break()
    
    # Technology Stack
    add_heading_with_style(doc, "7. Technology Stack", 1, RGBColor(0, 102, 204))
    
    tech_data = [
        ("Layer", "Technology", "Purpose"),
        ("Frontend", "Streamlit, Streamlit ACE", "Web UI framework"),
        ("Visualization", "Plotly, Pandas", "Charts and data visualization"),
        ("Code Analysis", "AST, Regex, Pattern Matching", "Static code analysis"),
        ("Security", "cryptography, hashlib", "Cryptographic operations"),
        ("ML/AI", "sklearn (implied), Custom ML", "Issue ranking & classification"),
        ("Report Generation", "ReportLab, python-docx", "PDF and Document generation"),
        ("AI Chat", "LLM Integration (Custom)", "Natural language processing"),
        ("Backend", "Python 3.x", "Core logic implementation")
    ]
    
    add_table_with_data(doc, tech_data[0], tech_data[1:])
    
    doc.add_page_break()
    
    # Module Organization
    add_heading_with_style(doc, "8. Project Structure", 1, RGBColor(0, 102, 204))
    
    structure = """
ai_bug_localizer/
├── app.py                           # Main Streamlit application
├── analyzer/                        # Security and code analysis modules
│   ├── static_analysis.py          # AST-based analysis
│   ├── runtime_analysis.py         # Runtime error detection
│   ├── owasp_analyzer.py           # OWASP Top 10 detection
│   ├── encryption_analyzer.py      # Crypto analysis
│   ├── compliance_checker.py       # GDPR, HIPAA, PCI-DSS
│   ├── authentication_analyzer.py  # Auth security review
│   ├── api_security_analyzer.py    # API security
│   ├── dependency_vulnerability_analyzer.py # Package vulnerabilities
│   ├── secure_configuration_analyzer.py    # Config security
│   ├── input_validation_analyzer.py        # Input validation
│   ├── threat_modeling_analyzer.py         # Threat modeling
│   ├── fix_suggester.py            # Fix recommendations
│   ├── code_corrector.py           # Auto-fix engine
│   ├── ranking.py                  # ML-based ranking
│   └── explanation.py              # Issue explanations
├── core/                            # Core processing engines
│   ├── ai_engine.py                # Main AI engine
│   ├── enhanced_ai_chat.py         # Advanced AI chat
│   ├── quality_engine.py           # Quality metrics
│   └── report_engine.py            # Report generation
├── ml/                              # Machine Learning
│   └── scoring.py                  # ML scoring system
├── ui/                              # UI Components
│   └── components.py               # Streamlit components
├── utils/                           # Utilities
│   └── pylint_runner.py            # Code linting
└── assets/                          # Static assets
    └── styles.css                  # Styling
    """
    
    for line in structure.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # Key Features Matrix
    add_heading_with_style(doc, "9. Feature Capability Matrix", 1, RGBColor(0, 102, 204))
    
    features_data = [
        ("Feature Category", "Capability", "Analyzer Module"),
        ("Code Analysis", "Static AST Analysis", "static_analysis.py"),
        ("Code Analysis", "Runtime Error Detection", "runtime_analysis.py"),
        ("Code Analysis", "Complexity Scoring", "complexity.py"),
        ("Security", "OWASP Top 10", "owasp_analyzer.py"),
        ("Security", "Encryption Validation", "encryption_analyzer.py"),
        ("Security", "Compliance (GDPR/HIPAA/PCI)", "compliance_checker.py"),
        ("Security", "Authentication Review", "authentication_analyzer.py"),
        ("Security", "API Security", "api_security_analyzer.py"),
        ("Security", "Dependency Vulnerabilities", "dependency_vulnerability_analyzer.py"),
        ("Security", "Configuration Security", "secure_configuration_analyzer.py"),
        ("Security", "Input Validation", "input_validation_analyzer.py"),
        ("Security", "Threat Modeling", "threat_modeling_analyzer.py"),
        ("Intelligence", "Issue Ranking (ML)", "ranking.py"),
        ("Intelligence", "Automated Fixes", "fix_suggester.py"),
        ("Intelligence", "AI Explanations", "explanation.py"),
        ("Intelligence", "Code Correction", "code_corrector.py"),
        ("Intelligence", "AI Chat Assistant", "enhanced_ai_chat.py"),
        ("Reporting", "Dashboard Visualization", "ui/components.py"),
        ("Reporting", "PDF Export", "core/report_engine.py"),
        ("Reporting", "Quality Metrics", "core/quality_engine.py")
    ]
    
    add_table_with_data(doc, features_data[0], features_data[1:])
    
    doc.add_page_break()
    
    # Processing Pipeline Summary
    add_heading_with_style(doc, "10. Complete Processing Pipeline", 1, RGBColor(0, 102, 204))
    
    pipeline_steps = [
        "User Input → Code submission via Streamlit UI",
        "Validation → Check code syntax and format",
        "Parsing → Generate AST and token stream",
        "Static Analysis → Pattern matching for bugs",
        "Runtime Analysis → Execute with error capture",
        "Security Scanning → Run all security modules",
        "Aggregation → Combine all findings",
        "Deduplication → Remove redundant issues",
        "Ranking → ML model assigns scores",
        "Classification → Categorize by severity",
        "Fix Generation → Generate fix suggestions",
        "Explanation → AI generates descriptions",
        "Report Build → Aggregate findings into report",
        "Visualization → Render charts and dashboards",
        "Export → Generate PDF report",
        "Display → Show results to user"
    ]
    
    for i, step in enumerate(pipeline_steps, 1):
        doc.add_paragraph(f"{i:2d}. {step}", style='List Number')
    
    doc.add_page_break()
    
    # Conclusion
    add_heading_with_style(doc, "Summary", 1, RGBColor(0, 102, 204))
    summary_text = """
The AI Bug Localizer PRO system is a comprehensive, layered architecture designed to provide 
enterprise-grade code analysis and security validation. The modular design allows for:

• Scalable addition of new security analyzers
• Independent testing and maintenance of components
• Flexible UI customization through Streamlit
• Real-time processing of code analysis
• Integration with external security tools and APIs

The combination of static analysis, runtime analysis, ML-based ranking, and security-specific modules 
creates a powerful tool for developers and security professionals to identify and remediate code issues 
and vulnerabilities efficiently.
    """
    doc.add_paragraph(summary_text)
    
    # Save document
    output_path = "AI_BUG_LOCALIZER_Architecture_Diagrams.docx"
    doc.save(output_path)
    print(f"✅ Document created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_word_document()
